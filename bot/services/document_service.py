"""
Сервис для генерации PDF и DOCX документов
"""
import io
from datetime import datetime
from typing import BinaryIO
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.lib.enums import TA_CENTER, TA_LEFT
from loguru import logger


class DocumentService:
    """Сервис для генерации документов"""
    
    @staticmethod
    def generate_docx(report_text: str, week_start: str, week_end: str) -> BinaryIO:
        """
        Сгенерировать DOCX документ с недельным отчетом
        
        Args:
            report_text: Содержимое отчета
            week_start: Дата начала недели
            week_end: Дата окончания недели
        
        Returns:
            BytesIO объект с содержимым DOCX
        """
        try:
            doc = Document()
            
            # Добавляем заголовок
            title = doc.add_heading('Еженедельный отчет / Həftəlik Hesabat', 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Добавляем диапазон дат
            date_para = doc.add_paragraph(f'Период / Dövr: {week_start} - {week_end}')
            date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            date_para.runs[0].font.size = Pt(12)
            
            doc.add_paragraph()  # Пустая строка
            
            # Добавляем содержимое отчета
            # Разбиваем по строкам и добавляем как параграфы
            for line in report_text.split('\n'):
                if line.strip():
                    para = doc.add_paragraph(line)
                    para.style.font.size = Pt(11)
            
            # Добавляем подвал
            doc.add_page_break()
            footer_para = doc.add_paragraph(
                f'Отчет сгенерирован автоматически / Hesabat avtomatik yaradılıb\n'
                f'{datetime.now().strftime("%d.%m.%Y %H:%M")}'
            )
            footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            footer_para.runs[0].font.size = Pt(9)
            footer_para.runs[0].italic = True
            
            # Сохраняем в BytesIO
            doc_io = io.BytesIO()
            doc.save(doc_io)
            doc_io.seek(0)
            
            logger.info("DOCX документ успешно сгенерирован")
            return doc_io
        
        except Exception as e:
            logger.error(f"Ошибка генерации DOCX: {e}")
            raise
    
    @staticmethod
    def generate_pdf(report_text: str, week_start: str, week_end: str) -> BinaryIO:
        """
        Сгенерировать PDF документ с недельным отчетом
        
        Args:
            report_text: Содержимое отчета
            week_start: Дата начала недели
            week_end: Дата окончания недели
        
        Returns:
            BytesIO объект с содержимым PDF
        """
        try:
            pdf_io = io.BytesIO()
            
            # Создаем PDF документ
            doc = SimpleDocTemplate(
                pdf_io,
                pagesize=A4,
                rightMargin=72,
                leftMargin=72,
                topMargin=72,
                bottomMargin=18
            )
            
            # Контейнер для элементов PDF
            story = []
            
            # Стили
            styles = getSampleStyleSheet()
            
            # Стиль заголовка
            title_style = ParagraphStyle(
                'CustomTitle',
                parent=styles['Heading1'],
                fontSize=18,
                textColor='#000000',
                spaceAfter=30,
                alignment=TA_CENTER,
                fontName='Helvetica-Bold'
            )
            
            # Обычный стиль
            normal_style = ParagraphStyle(
                'CustomNormal',
                parent=styles['Normal'],
                fontSize=11,
                textColor='#000000',
                spaceAfter=12,
                alignment=TA_LEFT,
                fontName='Helvetica'
            )
            
            # Стиль даты
            date_style = ParagraphStyle(
                'DateStyle',
                parent=styles['Normal'],
                fontSize=12,
                textColor='#666666',
                spaceAfter=20,
                alignment=TA_CENTER,
                fontName='Helvetica-Oblique'
            )
            
            # Добавляем заголовок
            title = Paragraph("Еженедельный отчет / Həftəlik Hesabat", title_style)
            story.append(title)
            
            # Добавляем диапазон дат
            date_text = f"Период / Dövr: {week_start} - {week_end}"
            date_para = Paragraph(date_text, date_style)
            story.append(date_para)
            
            story.append(Spacer(1, 0.2 * inch))
            
            # Добавляем содержимое отчета
            for line in report_text.split('\n'):
                if line.strip():
                    # Экранируем специальные символы для PDF
                    safe_line = line.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
                    para = Paragraph(safe_line, normal_style)
                    story.append(para)
            
            # Добавляем подвал
            story.append(PageBreak())
            footer_text = (
                f"Отчет сгенерирован автоматически / Hesabat avtomatik yaradılıb<br/>"
                f"{datetime.now().strftime('%d.%m.%Y %H:%M')}"
            )
            footer_style = ParagraphStyle(
                'FooterStyle',
                parent=styles['Normal'],
                fontSize=9,
                textColor='#999999',
                alignment=TA_CENTER,
                fontName='Helvetica-Oblique'
            )
            footer = Paragraph(footer_text, footer_style)
            story.append(footer)
            
            # Создаем PDF
            doc.build(story)
            
            pdf_io.seek(0)
            
            logger.info("PDF документ успешно сгенерирован")
            return pdf_io
        
        except Exception as e:
            logger.error(f"Ошибка генерации PDF: {e}")
            raise


# Глобальный экземпляр сервиса документов
document_service = DocumentService()